"""
Document Processing Utility

This module provides functions to extract text from various document formats
including PDF, DOCX, and TXT files for processing by SocialSynth-AI.
"""

import os
import logging
import tempfile
from typing import Optional, Union, BinaryIO, Any

# Configure logging
logger = logging.getLogger("socialsynth.utils.document_processor")

def extract_text_from_document(document_file: Union[str, BinaryIO]) -> str:
    """
    Extract text content from various document formats.
    
    Args:
        document_file: File path or file-like object (from Streamlit uploader)
        
    Returns:
        str: Extracted text content from the document
    """
    try:
        # Determine file type from name or content type
        if hasattr(document_file, 'name'):
            file_name = document_file.name
        elif isinstance(document_file, str):
            file_name = os.path.basename(document_file)
        else:
            logger.warning("Unable to determine file name, attempting detection by content")
            # Try to detect by first few bytes
            return _extract_text_by_content(document_file)
        
        # Process based on file extension
        file_ext = os.path.splitext(file_name)[1].lower()
        
        if file_ext == '.pdf':
            return _extract_text_from_pdf(document_file)
        elif file_ext == '.docx':
            return _extract_text_from_docx(document_file)
        elif file_ext == '.txt':
            return _extract_text_from_txt(document_file)
        else:
            logger.warning(f"Unsupported file format: {file_ext}")
            return "Unsupported document format. Please upload PDF, DOCX, or TXT files."
            
    except Exception as e:
        logger.error(f"Error extracting text from document: {e}")
        return f"Error processing document: {str(e)}"

def _extract_text_from_pdf(pdf_file: Union[str, BinaryIO]) -> str:
    """Extract text from PDF document"""
    try:
        import PyPDF2
        
        # If it's a file-like object from Streamlit, save to temp file first
        if not isinstance(pdf_file, str):
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(pdf_file.getvalue())
                temp_file_path = temp_file.name
            
            # Now process the temp file
            with open(temp_file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ""
                for page_num in range(len(pdf_reader.pages)):
                    text += pdf_reader.pages[page_num].extract_text() + "\n"
                
            # Clean up
            os.unlink(temp_file_path)
            return text
        else:
            # Direct file path
            with open(pdf_file, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ""
                for page_num in range(len(pdf_reader.pages)):
                    text += pdf_reader.pages[page_num].extract_text() + "\n"
            return text
            
    except ImportError:
        logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
        return "Error: PyPDF2 library not installed. Required for PDF processing."
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        return f"Error processing PDF document: {str(e)}"

def _extract_text_from_docx(docx_file: Union[str, BinaryIO]) -> str:
    """Extract text from DOCX document"""
    try:
        import docx
        
        # If it's a file-like object from Streamlit, save to temp file first
        if not isinstance(docx_file, str):
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(docx_file.getvalue())
                temp_file_path = temp_file.name
            
            # Now process the temp file
            doc = docx.Document(temp_file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Clean up
            os.unlink(temp_file_path)
            return text
        else:
            # Direct file path
            doc = docx.Document(docx_file)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        return "Error: python-docx library not installed. Required for DOCX processing."
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        return f"Error processing DOCX document: {str(e)}"

def _extract_text_from_txt(txt_file: Union[str, BinaryIO]) -> str:
    """Extract text from plain text document"""
    try:
        # If it's a file-like object from Streamlit
        if not isinstance(txt_file, str):
            return txt_file.getvalue().decode('utf-8')
        else:
            # Direct file path
            with open(txt_file, 'r', encoding='utf-8') as f:
                return f.read()
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {e}")
        return f"Error processing text document: {str(e)}"

def _extract_text_by_content(document_file: BinaryIO) -> str:
    """Try to detect file type by content and extract text"""
    try:
        # Read the first few bytes to try to determine file type
        header = document_file.read(5)
        document_file.seek(0)  # Reset to beginning
        
        # Check for PDF magic number (%PDF-)
        if header.startswith(b'%PDF-'):
            return _extract_text_from_pdf(document_file)
        # Check for DOCX (it's a zip file with specific content)
        elif header.startswith(b'PK\x03\x04'):
            return _extract_text_from_docx(document_file)
        else:
            # Try as plain text
            try:
                return document_file.getvalue().decode('utf-8')
            except UnicodeDecodeError:
                return "Unable to process this document format. Please use PDF, DOCX, or TXT files."
    except Exception as e:
        logger.error(f"Error detecting document type: {e}")
        return f"Error processing document: {str(e)}" 